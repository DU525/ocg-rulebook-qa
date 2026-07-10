import os
import re
import uuid
import asyncio
import aiohttp
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from .metadata_extractor import MetadataExtractor, DocumentType, ChunkMetadata

@dataclass
class DocumentChunk:
    """文档分块 data class"""
    id: str
    content: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    """文档处理器：解析、分块、向量化"""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.metadata_extractor = MetadataExtractor()

    def process_rst_file(self, file_path: str) -> List[DocumentChunk]:
        """处理RST格式规则文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        chunks = []
        sections = re.split(r'\n(?=#{1,3}\s)', content)

        current_chapter = ""
        current_section = ""

        for section in sections:
            lines = section.strip().split('\n')
            if not lines:
                continue

            title_line = lines[0]
            if title_line.startswith('#'):
                if len(lines[0]) >= 2 and lines[0][1] == '#':
                    current_section = lines[0].lstrip('#').strip()
                else:
                    current_chapter = lines[0].lstrip('#').strip()

            content_text = '\n'.join(lines[1:])
            text_chunks = self._split_text(content_text)

            for idx, chunk_text in enumerate(text_chunks):
                if len(chunk_text) < 50:
                    continue

                chunks.append(DocumentChunk(
                    id=f"{Path(file_path).stem}_{idx}",
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'chapter': current_chapter,
                        'section': current_section,
                        'chunk_index': idx
                    }
                ))

        return chunks

    def _split_text(self, text: str) -> List[str]:
        """智能文本分块，保留句子完整性"""
        chunks = []
        paragraphs = text.split('\n\n')
        current_chunk = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(para) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""

                sentences = re.split(r'([。！？；])', para)
                for i in range(0, len(sentences) - 1, 2):
                    sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else '')
                    if len(current_chunk) + len(sentence) > self.chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk)
                        current_chunk = sentence[self.chunk_overlap:]
                    else:
                        current_chunk += sentence
            else:
                if len(current_chunk) + len(para) > self.chunk_size:
                    chunks.append(current_chunk)
                    current_chunk = para
                else:
                    current_chunk += '\n\n' + para if current_chunk else para

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """处理PDF文档"""
        from pypdf import PdfReader

        chunks = []
        reader = PdfReader(file_path)

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text.strip():
                continue

            text_chunks = self._split_text(text)
            for idx, chunk_text in enumerate(text_chunks):
                chunks.append(DocumentChunk(
                    id=f"pdf_{page_num}_{idx}",
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'type': 'pdf',
                        'page': page_num + 1,
                        'chunk_index': idx
                    }
                ))

        return chunks

    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """处理DOCX文档"""
        from docx import Document

        chunks = []
        doc = Document(file_path)

        current_chapter = ""
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测标题
            if para.style.name.startswith('Heading'):
                if para.style.name == 'Heading 1':
                    current_chapter = text
                    current_section = ""
                elif para.style.name == 'Heading 2':
                    current_section = text
            else:
                text_chunks = self._split_text(text)
                for idx, chunk_text in enumerate(text_chunks):
                    if len(chunk_text) < 50:
                        continue
                    chunks.append(DocumentChunk(
                        id=f"docx_{hash(file_path)}_{idx}",
                        content=chunk_text,
                        metadata={
                            'source': str(file_path),
                            'chapter': current_chapter,
                            'section': current_section,
                            'chunk_index': idx
                        }
                    ))

        return chunks

    def process_txt(self, file_path: str) -> List[DocumentChunk]:
        """处理TXT文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        text_chunks = self._split_text(content)
        chunks = []

        for idx, chunk_text in enumerate(text_chunks):
            if len(chunk_text) < 50:
                continue
            chunks.append(DocumentChunk(
                id=f"txt_{hash(file_path)}_{idx}",
                content=chunk_text,
                metadata={
                    'source': str(file_path),
                    'chapter': '未分类',
                    'section': '',
                    'chunk_index': idx
                }
            ))

        return chunks

    async def fetch_url_content(self, url: str, selectors: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        异步抓取网页内容并解析
        
        Args:
            url: 目标网页URL
            selectors: 可选的CSS选择器列表，用于提取特定内容
            
        Returns:
            包含title, content, html的字典
        """
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as response:
                    if response.status != 200:
                        return {
                            'success': False,
                            'error': f'HTTP {response.status}: {response.reason}',
                            'url': url
                        }
                    
                    html = await response.text()
                    
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    title = soup.title.string if soup.title else url.split('/')[-1] or '未命名页面'
                    
                    if selectors:
                        content_parts = []
                        for selector in selectors:
                            elements = soup.select(selector)
                            content_parts.extend([elem.get_text(strip=True) for elem in elements if elem.get_text(strip=True)])
                        content = '\n\n'.join(content_parts) if content_parts else soup.get_text(strip=True)
                    else:
                        for tag in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                            tag.decompose()
                        content = soup.get_text(separator='\n', strip=True)
                    
                    content = re.sub(r'\n{3,}', '\n\n', content)
                    content = content.strip()
                    
                    return {
                        'success': True,
                        'url': url,
                        'title': title,
                        'content': content,
                        'content_length': len(content)
                    }
        except asyncio.TimeoutError:
            return {'success': False, 'error': '请求超时（30秒）', 'url': url}
        except Exception as e:
            return {'success': False, 'error': f'抓取失败: {str(e)}', 'url': url}

    def process_url_content(self, content: str, title: str, url: str) -> List[DocumentChunk]:
        """
        处理从URL抓取的内容，进行分块
        
        Args:
            content: 网页文本内容
            title: 网页标题
            url: 原始URL
            
        Returns:
            分块列表
        """
        text_chunks = self._split_text(content)
        chunks = []

        for idx, chunk_text in enumerate(text_chunks):
            if len(chunk_text) < 50:
                continue
            chunks.append(DocumentChunk(
                id=f"url_{uuid.uuid4().hex[:8]}_{idx}",
                content=chunk_text,
                metadata={
                    'source': url,
                    'title': title,
                    'type': 'url',
                    'chapter': title,
                    'section': '',
                    'chunk_index': idx
                }
            ))

        return chunks

    def process_markdown_file(self, file_path: str) -> List[DocumentChunk]:
        """处理 Markdown 格式文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        chunks = []
        doc_type = DocumentType.MARKDOWN
        
        # 按标题分割内容
        sections = re.split(r'\n(?=#{1,6}\s)', content)
        
        current_chapter = ""
        current_section = ""
        current_subsection = ""
        parent_stack = []
        heading_level_stack = []
        
        for section in sections:
            lines = section.strip().split('\n')
            if not lines:
                continue

            # 提取标题信息
            heading, level = self.metadata_extractor.extract_heading_info(section, doc_type)
            
            if heading:
                # 管理父块栈
                while heading_level_stack and heading_level_stack[-1] >= level:
                    heading_level_stack.pop()
                    if parent_stack:
                        parent_stack.pop()
                
                if level == 1:
                    current_chapter = heading
                    current_section = ""
                    current_subsection = ""
                elif level == 2:
                    current_section = heading
                    current_subsection = ""
                elif level == 3:
                    current_subsection = heading
                
                # 当前标题块作为父块
                parent_id = f"{Path(file_path).stem}_{len(chunks)}"
                parent_stack.append(parent_id)
                heading_level_stack.append(level)
                
                # 添加标题块
                chunks.append(DocumentChunk(
                    id=parent_id,
                    content=heading,
                    metadata={
                        'source': str(file_path),
                        'type': 'markdown',
                        'chapter': current_chapter,
                        'section': current_section,
                        'subsection': current_subsection,
                        'level': level,
                        'heading': heading,
                        'chunk_index': len(chunks),
                        'is_heading': True
                    }
                ))

            # 处理内容
            content_text = '\n'.join(lines[1:]) if heading else section
            text_chunks = self._split_text(content_text)
            
            parent_id = parent_stack[-1] if parent_stack else None
            
            for idx, chunk_text in enumerate(text_chunks):
                if len(chunk_text) < 50:
                    continue

                chunks.append(DocumentChunk(
                    id=f"{Path(file_path).stem}_{len(chunks)}",
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'type': 'markdown',
                        'chapter': current_chapter,
                        'section': current_section,
                        'subsection': current_subsection,
                        'level': level if heading else (heading_level_stack[-1] if heading_level_stack else 0),
                        'chunk_index': len(chunks),
                        'parent_id': parent_id,
                        'is_heading': False
                    }
                ))

        return chunks

    def process_with_metadata_inheritance(self, file_path: str) -> List[DocumentChunk]:
        """
        处理文档并支持父子块元数据继承
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            DocumentChunk 列表，包含完整的元数据和父子关系
        """
        doc_type = self.metadata_extractor.identify_document_type(file_path)
        
        if doc_type == DocumentType.RST:
            return self._process_rst_with_inheritance(file_path)
        elif doc_type == DocumentType.MARKDOWN:
            return self.process_markdown_file(file_path)
        elif doc_type == DocumentType.PDF:
            return self._process_pdf_with_inheritance(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._process_docx_with_inheritance(file_path)
        elif doc_type == DocumentType.TXT:
            return self._process_txt_with_inheritance(file_path)
        else:
            return self.process_txt(file_path)

    def _process_rst_with_inheritance(self, file_path: str) -> List[DocumentChunk]:
        """处理 RST 文档，支持父子块元数据继承"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        chunks = []
        doc_type = DocumentType.RST
        
        # 按标题分割内容
        sections = re.split(r'\n(?=.+?\n[=]+)', content)
        
        current_chapter = ""
        current_section = ""
        parent_stack = []
        heading_level_stack = []
        
        for section in sections:
            lines = section.strip().split('\n')
            if not lines:
                continue

            # 提取标题信息
            heading, level = self.metadata_extractor.extract_heading_info(section, doc_type)
            
            if heading:
                # 管理父块栈
                while heading_level_stack and heading_level_stack[-1] >= level:
                    heading_level_stack.pop()
                    if parent_stack:
                        parent_stack.pop()
                
                if level == 1:
                    current_chapter = heading
                    current_section = ""
                elif level == 2:
                    current_section = heading
                
                # 当前标题块作为父块
                parent_id = f"{Path(file_path).stem}_{len(chunks)}"
                parent_stack.append(parent_id)
                heading_level_stack.append(level)
                
                # 添加标题块
                chunks.append(DocumentChunk(
                    id=parent_id,
                    content=heading,
                    metadata={
                        'source': str(file_path),
                        'type': 'rst',
                        'chapter': current_chapter,
                        'section': current_section,
                        'level': level,
                        'heading': heading,
                        'chunk_index': len(chunks),
                        'is_heading': True
                    }
                ))

            # 处理内容
            content_start = 2 if heading else 0
            content_text = '\n'.join(lines[content_start:])
            text_chunks = self._split_text(content_text)
            
            parent_id = parent_stack[-1] if parent_stack else None
            
            for idx, chunk_text in enumerate(text_chunks):
                if len(chunk_text) < 50:
                    continue

                chunks.append(DocumentChunk(
                    id=f"{Path(file_path).stem}_{len(chunks)}",
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'type': 'rst',
                        'chapter': current_chapter,
                        'section': current_section,
                        'level': level if heading else (heading_level_stack[-1] if heading_level_stack else 0),
                        'chunk_index': len(chunks),
                        'parent_id': parent_id,
                        'is_heading': False
                    }
                ))

        return chunks

    def _process_pdf_with_inheritance(self, file_path: str) -> List[DocumentChunk]:
        """处理 PDF 文档，支持父子块元数据继承"""
        from pypdf import PdfReader

        chunks = []
        reader = PdfReader(file_path)
        parent_stack = []
        page_parent_map = {}

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text.strip():
                continue

            # 尝试从页面开头提取标题
            lines = text.strip().split('\n')
            potential_heading = None
            if lines:
                first_line = lines[0].strip()
                # 检查是否是可能的标题
                if len(first_line) < 200 and not re.match(r'^\d', first_line):
                    potential_heading = first_line

            if potential_heading:
                parent_id = f"pdf_page_{page_num}_heading"
                parent_stack = [parent_id]
                page_parent_map[page_num] = parent_id
                
                # 添加标题块
                chunks.append(DocumentChunk(
                    id=parent_id,
                    content=potential_heading,
                    metadata={
                        'source': str(file_path),
                        'type': 'pdf',
                        'page': page_num + 1,
                        'heading': potential_heading,
                        'chunk_index': len(chunks),
                        'is_heading': True
                    }
                ))

            text_chunks = self._split_text(text)
            parent_id = parent_stack[-1] if parent_stack else None
            
            for idx, chunk_text in enumerate(text_chunks):
                chunks.append(DocumentChunk(
                    id=f"pdf_{page_num}_{len(chunks)}",
                    content=chunk_text,
                    metadata={
                        'source': str(file_path),
                        'type': 'pdf',
                        'page': page_num + 1,
                        'chunk_index': len(chunks),
                        'parent_id': parent_id,
                        'is_heading': False
                    }
                ))

        return chunks

    def _process_docx_with_inheritance(self, file_path: str) -> List[DocumentChunk]:
        """处理 DOCX 文档，支持父子块元数据继承"""
        from docx import Document

        chunks = []
        doc = Document(file_path)
        parent_stack = []
        heading_level_stack = []
        current_chapter = ""
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if len(para.style.name.split()) > 1 else 1
                
                # 管理父块栈
                while heading_level_stack and heading_level_stack[-1] >= level:
                    heading_level_stack.pop()
                    if parent_stack:
                        parent_stack.pop()
                
                if level == 1:
                    current_chapter = text
                    current_section = ""
                elif level == 2:
                    current_section = text
                
                # 当前标题块作为父块
                parent_id = f"docx_heading_{len(chunks)}"
                parent_stack.append(parent_id)
                heading_level_stack.append(level)
                
                # 添加标题块
                chunks.append(DocumentChunk(
                    id=parent_id,
                    content=text,
                    metadata={
                        'source': str(file_path),
                        'type': 'docx',
                        'chapter': current_chapter,
                        'section': current_section,
                        'level': level,
                        'heading': text,
                        'chunk_index': len(chunks),
                        'is_heading': True
                    }
                ))
            else:
                text_chunks = self._split_text(text)
                parent_id = parent_stack[-1] if parent_stack else None
                
                for idx, chunk_text in enumerate(text_chunks):
                    if len(chunk_text) < 50:
                        continue
                    chunks.append(DocumentChunk(
                        id=f"docx_{hash(file_path)}_{len(chunks)}",
                        content=chunk_text,
                        metadata={
                            'source': str(file_path),
                            'type': 'docx',
                            'chapter': current_chapter,
                            'section': current_section,
                            'level': heading_level_stack[-1] if heading_level_stack else 0,
                            'chunk_index': len(chunks),
                            'parent_id': parent_id,
                            'is_heading': False
                        }
                    ))

        return chunks

    def _process_txt_with_inheritance(self, file_path: str) -> List[DocumentChunk]:
        """处理 TXT 文档，支持父子块元数据继承"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        chunks = []
        doc_type = DocumentType.TXT
        paragraphs = content.split('\n\n')
        
        parent_stack = []
        heading_level_stack = []
        current_chapter = ""
        current_section = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # 尝试提取标题
            heading, level = self.metadata_extractor.extract_heading_info(para, doc_type)
            
            if heading:
                # 管理父块栈
                while heading_level_stack and heading_level_stack[-1] >= level:
                    heading_level_stack.pop()
                    if parent_stack:
                        parent_stack.pop()
                
                if level == 1:
                    current_chapter = heading
                    current_section = ""
                elif level == 2:
                    current_section = heading
                
                # 当前标题块作为父块
                parent_id = f"txt_heading_{len(chunks)}"
                parent_stack.append(parent_id)
                heading_level_stack.append(level)
                
                # 添加标题块
                chunks.append(DocumentChunk(
                    id=parent_id,
                    content=heading,
                    metadata={
                        'source': str(file_path),
                        'type': 'txt',
                        'chapter': current_chapter,
                        'section': current_section,
                        'level': level,
                        'heading': heading,
                        'chunk_index': len(chunks),
                        'is_heading': True
                    }
                ))
            else:
                text_chunks = self._split_text(para)
                parent_id = parent_stack[-1] if parent_stack else None
                
                for idx, chunk_text in enumerate(text_chunks):
                    if len(chunk_text) < 50:
                        continue
                    chunks.append(DocumentChunk(
                        id=f"txt_{hash(file_path)}_{len(chunks)}",
                        content=chunk_text,
                        metadata={
                            'source': str(file_path),
                            'type': 'txt',
                            'chapter': current_chapter,
                            'section': current_section,
                            'level': heading_level_stack[-1] if heading_level_stack else 0,
                            'chunk_index': len(chunks),
                            'parent_id': parent_id,
                            'is_heading': False
                        }
                    ))

        return chunks