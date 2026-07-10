import os
import re
from typing import List, Dict, Any, Optional, Set
from dataclasses import dataclass
from pathlib import Path


@dataclass
class CleanResult:
    """清理结果类"""
    original_text: str
    cleaned_text: str
    removed_header: Optional[str] = None
    removed_footer: Optional[str] = None
    removed_lines: List[str] = None
    source: str = ""


class DocumentCleaner:
    """文档清理器：检测和清理页眉页脚"""

    def __init__(self):
        self.header_patterns = self._get_common_header_patterns()
        self.footer_patterns = self._get_common_footer_patterns()

    def _get_common_header_patterns(self) -> List[str]:
        """
        获取常见页眉模式
        
        Returns:
            正则表达式模式列表
        """
        return [
            r'^第\s*\d+\s*页$',
            r'^Page\s*\d+$',
            r'^page\s*\d+$',
            r'^-?\s*\d+\s*-?$',
            r'^\d+$',
            r'^www\..+\.com$',
            r'^http[s]?://',
            r'^Confidential$',
            r'^机密$',
            r'^版权所有$',
            r'^Copyright.*$',
            r'^Document.*$',
            r'^文档.*$',
            r'^\d+\s*-\s*\d+$',
            r'^[A-Za-z]+.*\d+$',
        ]

    def _get_common_footer_patterns(self) -> List[str]:
        """
        获取常见页脚模式
        
        Returns:
            正则表达式模式列表
        """
        return [
            r'^第\s*\d+\s*页$',
            r'^Page\s*\d+$',
            r'^page\s*\d+$',
            r'^-?\s*\d+\s*-?$',
            r'^\d+$',
            r'^www\..+\.com$',
            r'^http[s]?://',
            r'^Confidential$',
            r'^机密$',
            r'^版权所有$',
            r'^Copyright.*$',
            r'^Document.*$',
            r'^文档.*$',
            r'^\d+\s*-\s*\d+$',
            r'^[A-Za-z]+.*\d+$',
            r'^打印时间.*$',
            r'^生成时间.*$',
            r'^Created.*$',
            r'^Last modified.*$',
        ]

    def clean_text(self, text: str, source: str = "") -> CleanResult:
        """
        清理文本，移除页眉页脚
        
        Args:
            text: 原始文本
            source: 来源标识
            
        Returns:
            CleanResult对象
        """
        lines = text.split('\n')

        if len(lines) < 3:
            return CleanResult(
                original_text=text,
                cleaned_text=text,
                removed_lines=[],
                source=source
            )

        # 查找重复的页眉和页脚
        header_candidates = self._find_repeating_headers(lines)
        footer_candidates = self._find_repeating_footers(lines)

        cleaned_lines = []
        removed_lines = []
        removed_header = None
        removed_footer = None

        for line_num, line in enumerate(lines):
            line_stripped = line.strip()

            is_header = False
            is_footer = False

            # 检查是否是页眉
            if line_num in [0, 1] and header_candidates:
                for candidate in header_candidates:
                    if self._line_matches_pattern(line_stripped, candidate):
                        is_header = True
                        if not removed_header:
                            removed_header = line_stripped
                        break

            # 检查是否是页脚
            if line_num in [len(lines) - 1, len(lines) - 2] and footer_candidates:
                for candidate in footer_candidates:
                    if self._line_matches_pattern(line_stripped, candidate):
                        is_footer = True
                        if not removed_footer:
                            removed_footer = line_stripped
                        break

            # 检查常见模式
            if not is_header and not is_footer:
                if self._matches_common_patterns(line_stripped):
                    if line_num < len(lines) // 4:
                        is_header = True
                        if not removed_header:
                            removed_header = line_stripped
                    elif line_num > len(lines) * 3 // 4:
                        is_footer = True
                        if not removed_footer:
                            removed_footer = line_stripped

            if is_header or is_footer:
                removed_lines.append(line)
            else:
                cleaned_lines.append(line)

        cleaned_text = '\n'.join(cleaned_lines)

        return CleanResult(
            original_text=text,
            cleaned_text=cleaned_text,
            removed_header=removed_header,
            removed_footer=removed_footer,
            removed_lines=removed_lines,
            source=source
        )

    def _find_repeating_headers(self, lines: List[str]) -> Set[str]:
        """
        查找重复出现的页眉候选
        
        Args:
            lines: 文本行列表
            
        Returns:
            页眉候选集合
        """
        # 查看每个页面顶部的几行
        header_candidates = {}
        page_size_estimate = self._estimate_page_size(lines)

        for i in range(0, len(lines), max(page_size_estimate, 10)):
            for j in range(min(3, len(lines) - i)):
                line = lines[i + j].strip()
                if line and len(line) < 100:
                    normalized_line = self._normalize_line(line)
                    if normalized_line:
                        header_candidates[normalized_line] = header_candidates.get(normalized_line, 0) + 1

        # 保留出现多次的候选
        return {line for line, count in header_candidates.items() if count >= 2}

    def _find_repeating_footers(self, lines: List[str]) -> Set[str]:
        """
        查找重复出现的页脚候选
        
        Args:
            lines: 文本行列表
            
        Returns:
            页脚候选集合
        """
        footer_candidates = {}
        page_size_estimate = self._estimate_page_size(lines)

        for i in range(page_size_estimate - 1, len(lines), max(page_size_estimate, 10)):
            for j in range(min(3, i + 1)):
                if i - j >= 0 and i - j < len(lines):
                    line = lines[i - j].strip()
                    if line and len(line) < 100:
                        normalized_line = self._normalize_line(line)
                        if normalized_line:
                            footer_candidates[normalized_line] = footer_candidates.get(normalized_line, 0) + 1

        return {line for line, count in footer_candidates.items() if count >= 2}

    def _estimate_page_size(self, lines: List[str]) -> int:
        """
        估计每页的行数
        
        Args:
            lines: 文本行列表
            
        Returns:
            估计的每页行数
        """
        # 查找页码模式
        page_markers = []
        for i, line in enumerate(lines):
            if re.match(r'^-?\s*\d+\s*-?$', line.strip()) or re.match(r'^Page\s*\d+$', line.strip(), re.IGNORECASE):
                page_markers.append(i)

        if len(page_markers) >= 2:
            distances = [page_markers[i] - page_markers[i - 1] for i in range(1, len(page_markers))]
            if distances:
                return int(sum(distances) / len(distances))

        # 默认估计
        return 50

    def _normalize_line(self, line: str) -> str:
        """
        规范化文本行，用于比较
        
        Args:
            line: 文本行
            
        Returns:
            规范化后的文本
        """
        # 移除数字
        normalized = re.sub(r'\d+', '##', line)
        # 移除日期
        normalized = re.sub(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}', '##-##-##', normalized)
        # 移除时间
        normalized = re.sub(r'\d{1,2}:\d{2}(:\d{2})?', '##:##', normalized)
        # 移除特殊字符并简化
        normalized = re.sub(r'[^\w\s]', '', normalized)
        return normalized.strip().lower()

    def _line_matches_pattern(self, line: str, pattern: str) -> bool:
        """
        检查行是否匹配模式
        
        Args:
            line: 文本行
            pattern: 模式（规范化后的）
            
        Returns:
            是否匹配
        """
        normalized_line = self._normalize_line(line)
        return pattern in normalized_line or normalized_line in pattern

    def _matches_common_patterns(self, line: str) -> bool:
        """
        检查行是否匹配常见的页眉页脚模式
        
        Args:
            line: 文本行
            
        Returns:
            是否匹配
        """
        all_patterns = self.header_patterns + self.footer_patterns
        for pattern in all_patterns:
            if re.match(pattern, line.strip()):
                return True
        return False

    def clean_pdf_text(self, pdf_path: str) -> List[CleanResult]:
        """
        清理PDF文档的文本
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            CleanResult列表（每页一个）
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(pdf_path)
            results = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                result = self.clean_text(text, source=f"{pdf_path} (page {page_num + 1})")
                results.append(result)

            return results
        except Exception as e:
            print(f"PDF清理失败: {str(e)}")
            return []

    def clean_docx_text(self, docx_path: str) -> List[CleanResult]:
        """
        清理Word文档的文本
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            CleanResult列表
        """
        try:
            from docx import Document

            doc = Document(docx_path)
            full_text = []

            # 提取正文
            for para in doc.paragraphs:
                full_text.append(para.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            combined_text = '\n'.join(full_text)
            result = self.clean_text(combined_text, source=docx_path)

            return [result]
        except Exception as e:
            print(f"Word文档清理失败: {str(e)}")
            return []

    def clean_file(self, file_path: str) -> Dict[str, Any]:
        """
        根据文件类型清理文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            清理结果字典
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            results = self.clean_pdf_text(file_path)
            total_removed = sum(len(r.removed_lines) for r in results)
            combined_cleaned = '\n\n'.join(r.cleaned_text for r in results)

            return {
                'success': True,
                'type': 'pdf',
                'results': results,
                'total_removed_lines': total_removed,
                'cleaned_text': combined_cleaned,
                'source': file_path
            }
        elif file_ext in ['.docx', '.doc']:
            results = self.clean_docx_text(file_path)
            total_removed = sum(len(r.removed_lines) for r in results) if results else 0

            return {
                'success': True,
                'type': 'docx',
                'results': results,
                'total_removed_lines': total_removed,
                'cleaned_text': results[0].cleaned_text if results else '',
                'source': file_path
            }
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            result = self.clean_text(text, source=file_path)

            return {
                'success': True,
                'type': 'txt',
                'results': [result],
                'total_removed_lines': len(result.removed_lines),
                'cleaned_text': result.cleaned_text,
                'source': file_path
            }
        else:
            return {
                'success': False,
                'error': f'不支持的文件格式: {file_ext}',
                'type': 'unknown',
                'source': file_path
            }

    def add_custom_header_pattern(self, pattern: str) -> None:
        """
        添加自定义页眉模式
        
        Args:
            pattern: 正则表达式模式
        """
        if pattern not in self.header_patterns:
            self.header_patterns.append(pattern)

    def add_custom_footer_pattern(self, pattern: str) -> None:
        """
        添加自定义页脚模式
        
        Args:
            pattern: 正则表达式模式
        """
        if pattern not in self.footer_patterns:
            self.footer_patterns.append(pattern)
