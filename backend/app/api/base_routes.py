"""Shared base route handlers for OCG and DM domains."""
from flask import Blueprint, request, jsonify, Response, stream_with_context
from app.services.document_processor import DocumentProcessor
from app.services.feedback_analyzer import FeedbackAnalyzer
from app.services.feishu_notifier import send_alert as feishu_send_alert, check_feishu_cli_available
from app.services.conversation_memory import get_conversation_memory
from app.config import Config
from collections import Counter
from typing import Dict, Any, Optional
import uuid
import logging
import time
import threading
import os
import re
import json
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)

upload_tasks = {}
MAX_FILE_SIZE = 50 * 1024 * 1024


FEEDBACK_REASON_CATEGORIES = [
    "answer_inaccurate",
    "answer_irrelevant",
    "citation_missing",
    "format_issue",
    "outdated_info",
    "other"
]

class BaseRouteHandler:
    """Shared base route handler for both OCG and DM domains."""

    def __init__(self, blueprint: Blueprint, vector_store, db, rag_engine,
                 game_type: str, prompt_templates: dict, models: dict,
                 agent_instance=None, skill_registry=None):
        self.bp = blueprint
        self.vector_store = vector_store
        self.db = db
        self.rag_engine = rag_engine
        self.game_type = game_type
        self.prompt_templates = prompt_templates
        self.Conversation = models['Conversation']
        self.Message = models['Message']
        self.Document = models['Document']
        self.Alert = models['Alert']
        self.AlertRule = models['AlertRule']
        self.PerformanceLog = models['PerformanceLog']
        self.Feedback = models['Feedback']
        self.NegativeSample = models.get('NegativeSample')
        self.agent_instance = agent_instance
        self.skill_registry = skill_registry
        self.rag_runtime_config = {
            'top_k': 5,
            'temperature': 0.3,
            'max_tokens': 1500,
            'system_prompt_template': 'default',
            'streaming_enabled': False,
            'similarity_threshold': 0.5,
        }
        self._register_routes()

    def _register_routes(self):
        bp = self.bp
        bp.route('/health', methods=['GET'])(self.health_check)
        bp.route('/chat/question', methods=['POST'])(self.ask_question)
        bp.route('/chat/question/stream', methods=['POST'])(self.ask_question_stream)
        bp.route('/conversations', methods=['GET'])(self.get_conversations)
        bp.route('/conversations/<conversation_id>', methods=['GET'])(self.get_conversation)
        bp.route('/conversations/<conversation_id>', methods=['DELETE'])(self.delete_conversation)
        bp.route('/documents', methods=['GET'])(self.get_documents)
        bp.route('/documents/search', methods=['GET'])(self.search_documents)
        bp.route('/documents/upload', methods=['POST'])(self.upload_document)
        bp.route('/documents/<doc_id>/progress', methods=['GET'])(self.get_upload_progress)
        bp.route('/documents/<doc_id>', methods=['DELETE'])(self.delete_document)
        bp.route('/documents/<doc_id>/preview', methods=['GET'])(self.get_document_preview)
        bp.route('/documents/<doc_id>/chunks', methods=['GET'])(self.get_document_chunks)
        bp.route('/documents/import-url', methods=['POST'])(self.import_url)
        bp.route('/documents/<doc_id>/import-status', methods=['GET'])(self.get_import_status)
        bp.route('/conversations/search', methods=['GET'])(self.search_conversations)
        bp.route('/citations/<chunk_id>/detail', methods=['GET'])(self.get_citation_detail)
        bp.route('/config/models', methods=['GET'])(self.get_available_models)
        bp.route('/config/models', methods=['PUT'])(self.switch_model)
        bp.route('/feedback', methods=['POST'])(self.submit_feedback)
        bp.route('/feedback/<message_id>', methods=['GET'])(self.get_feedback_status)
        bp.route('/feedback/stats', methods=['GET'])(self.get_feedback_stats)
        bp.route('/feedback/negative-samples', methods=['GET'])(self.get_negative_samples)
        bp.route('/metrics', methods=['GET'])(self.get_metrics)
        bp.route('/metrics/alerts', methods=['GET'])(self.get_alert_rules)
        bp.route('/metrics/alerts', methods=['POST'])(self.create_or_update_alert_rule)
        bp.route('/metrics/alerts/history', methods=['GET'])(self.get_alert_history)
        bp.route('/metrics/alerts/<alert_id>/read', methods=['PUT'])(self.mark_alert_read)
        bp.route('/metrics/alerts/unread', methods=['GET'])(self.get_unread_alerts)
        bp.route('/metrics/trend', methods=['GET'])(self.get_metrics_trend)
        bp.route('/metrics/search-quality', methods=['GET'])(self.get_search_quality)
        bp.route('/metrics/quality', methods=['GET'])(self.get_quality_metrics)
        bp.route('/metrics/cache', methods=['GET'])(self.get_cache_stats)
        bp.route('/config/rag', methods=['GET'])(self.get_rag_config)
        bp.route('/config/rag', methods=['PUT'])(self.update_rag_config)
        if self.agent_instance:
            bp.route('/chat/agent', methods=['POST'])(self.agent_query)
            bp.route('/agent/skills', methods=['GET'])(self.list_skills)
        if self.rag_engine and hasattr(self.rag_engine, 'query_multimodal'):
            bp.route('/chat/multimodal', methods=['POST'])(self.multimodal_query)

    def health_check(self):
        stats = self.vector_store.get_collection_stats() if self.vector_store else {}
        return jsonify({
            'status': 'healthy',
            'vector_store_count': stats.get('count', 0),
            'timestamp': datetime.utcnow().isoformat()
        })

    def _get_system_prompt(self):
        template_id = self.rag_runtime_config['system_prompt_template']
        return self.prompt_templates.get(template_id, self.prompt_templates.get('default', ''))

    def _resolve_search_path(self, requested_search_type: str = None) -> Dict[str, Any]:
        """根据请求和配置决定检索路径
        
        Args:
            requested_search_type: 请求体里的 search_type，可选 'hybrid' / 'vector' / 'bm25' / None
            
        Returns:
            Dict 包含:
                - search_type: 实际生效的检索类型（'hybrid' / 'vector'）
                - use_multi_stage: 是否启用多阶段 RRF 融合
                - path_label: 给前端展示的路径标签
                - rrf_enabled: 全局 RRF 开关
        """
        rrf_enabled = bool(getattr(Config, 'RRF_ENABLED', True))
        requested = (requested_search_type or '').lower().strip()

        if requested == 'hybrid' and rrf_enabled:
            return {
                'search_type': 'hybrid',
                'use_multi_stage': True,
                'path_label': 'BM25+向量+RRF+Cross-Encoder 多阶段融合',
                'rrf_enabled': True,
                'requested': requested or 'default(hybrid)',
            }
        if requested == 'vector':
            return {
                'search_type': 'vector',
                'use_multi_stage': False,
                'path_label': '纯 FAISS 向量检索',
                'rrf_enabled': rrf_enabled,
                'requested': requested,
            }
        if requested == 'hybrid' and not rrf_enabled:
            # 显式请求 hybrid 但全局 RRF 关闭 → 回退纯向量
            return {
                'search_type': 'vector',
                'use_multi_stage': False,
                'path_label': 'RRF 全局关闭，回退纯 FAISS 向量检索',
                'rrf_enabled': False,
                'requested': 'hybrid(fallback)',
            }
        # 默认：未指定时，遵循全局 RRF_ENABLED
        if rrf_enabled:
            return {
                'search_type': 'hybrid',
                'use_multi_stage': True,
                'path_label': '默认: BM25+向量+RRF+Cross-Encoder 多阶段融合',
                'rrf_enabled': True,
                'requested': 'default',
            }
        return {
            'search_type': 'vector',
            'use_multi_stage': False,
            'path_label': '默认: 纯 FAISS 向量检索（RRF 关闭）',
            'rrf_enabled': False,
            'requested': 'default',
        }

    def ask_question(self):
        start_time = time.time()
        data = request.get_json(silent=True) or {}
        question = data.get('question', '')
        conversation_id = data.get('conversation_id')
        requested_search_type = data.get('search_type')

        if not question:
            return jsonify({'success': False, 'error': {'code': 'EMPTY_QUESTION', 'message': '问题不能为空'}}), 400

        # 决定检索路径（hybrid / vector）
        search_meta = self._resolve_search_path(requested_search_type)
        use_multi_stage = search_meta['use_multi_stage']

        session = self.db.get_session()
        try:
            if not conversation_id:
                conversation_id = str(uuid.uuid4())
                conversation = self.Conversation(id=conversation_id, title=question[:50])
                session.add(conversation)
                session.flush()

            messages = session.query(self.Message).filter_by(conversation_id=conversation_id).order_by(self.Message.created_at).all()
            history = [{'role': msg.role, 'content': msg.content} for msg in messages[-6:]]

            user_msg = self.Message(id=str(uuid.uuid4()), conversation_id=conversation_id, role='user', content=question)
            session.add(user_msg)
            session.commit()
        finally:
            session.close()

        conversation_memory = get_conversation_memory()
        conversation_memory.add_message(conversation_id, 'user', question)
        memory_history = conversation_memory.get_history(conversation_id, max_turns=5)
        if memory_history:
            history = memory_history

        system_prompt = self._get_system_prompt()
        # 把 per-request 的 use_multi_stage 真正透传给 RAGEngine
        # 这样 search_type=vector 时真走纯向量，search_type=hybrid 时真走 RRF
        response = self.rag_engine.query(
            question,
            top_k=self.rag_runtime_config['top_k'],
            conversation_history=history,
            temperature=self.rag_runtime_config['temperature'],
            max_tokens=self.rag_runtime_config['max_tokens'],
            system_prompt=system_prompt,
            use_multi_stage=use_multi_stage,
        )

        conversation_memory.add_message(conversation_id, 'assistant', response.answer)

        session2 = self.db.get_session()
        try:
            assistant_msg = self.Message(id=str(uuid.uuid4()), conversation_id=conversation_id,
                role='assistant', content=response.answer, citations=response.citations)
            session2.add(assistant_msg)
            session2.commit()
        finally:
            session2.close()

        response_time = int((time.time() - start_time) * 1000)
        session3 = self.db.get_session()
        try:
            log = self.PerformanceLog(id=str(uuid.uuid4()), endpoint=self.bp.url_prefix + '/chat/question',
                method='POST', response_time_ms=response_time, status_code=200, tokens_used=0, created_at=datetime.utcnow())
            session3.add(log)
            session3.commit()
        except Exception as e:
            logger.warning(f"Performance log error: {e}")
        finally:
            session3.close()

        return jsonify({'success': True, 'data': {
            'answer': response.answer,
            'citations': response.citations,
            'confidence': response.confidence,
            'conversation_id': conversation_id,
            'response_time_ms': response_time,
            'search_type': search_meta['search_type'],
            'search_path': search_meta['path_label'],
            'rrf_enabled': search_meta['rrf_enabled'],
            'requested_search_type': search_meta['requested'],
        }})

    def ask_question_stream(self):
        data = request.get_json(silent=True) or {}
        question = data.get('question', '')
        conversation_id = data.get('conversation_id')
        requested_search_type = data.get('search_type')

        if not question:
            return jsonify({'success': False, 'error': {'code': 'EMPTY_QUESTION', 'message': '问题不能为空'}}), 400

        # 决定检索路径（hybrid / vector）
        search_meta = self._resolve_search_path(requested_search_type)
        use_multi_stage = search_meta['use_multi_stage']

        def generate():
            import queue
            import threading

            full_answer = []
            citations_data = None
            confidence_data = 0.0
            final_conversation_id = conversation_id
            system_prompt = self._get_system_prompt()

            try:
                yield "data: {\"status\":\"stream_start\"}\n\n"

                session = self.db.get_session()
                try:
                    if not final_conversation_id:
                        final_conversation_id = str(uuid.uuid4())
                        conversation = self.Conversation(id=final_conversation_id, title=question[:50])
                        session.add(conversation)
                        session.flush()

                    messages = session.query(self.Message).filter_by(conversation_id=final_conversation_id).order_by(self.Message.created_at).all()
                    history = [{'role': msg.role, 'content': msg.content} for msg in messages[-6:]]

                    user_msg = self.Message(id=str(uuid.uuid4()), conversation_id=final_conversation_id, role='user', content=question)
                    session.add(user_msg)
                    session.commit()
                finally:
                    session.close()

                for chunk in self.rag_engine.query_stream(question, top_k=self.rag_runtime_config['top_k'],
                    conversation_history=history, temperature=self.rag_runtime_config['temperature'],
                    max_tokens=self.rag_runtime_config['max_tokens'], system_prompt=system_prompt,
                    use_multi_stage=use_multi_stage):
                    full_answer.append(chunk)
                    yield f"data: {chunk}\n\n"

                full_text = ""
                for c in full_answer:
                    try:
                        decoded = json.loads(c.replace("data: ", ""))
                        if "content" in decoded:
                            full_text += decoded["content"]
                        if "citations" in decoded:
                            citations_data = decoded["citations"]
                            confidence_data = decoded.get("confidence", 0.0)
                    except:
                        pass
                    if "data: [DONE]" in c:
                        break

                if citations_data:
                    yield f"data: {json.dumps({'citations': citations_data, 'confidence': confidence_data, 'conversation_id': final_conversation_id})}\n\n"

                yield "data: [DONE]\n\n"

                if full_text:
                    def save_to_db():
                        try:
                            session2 = self.db.get_session()
                            try:
                                assistant_msg = self.Message(id=str(uuid.uuid4()), conversation_id=final_conversation_id,
                                    role='assistant', content=full_text, citations=citations_data or [])
                                session2.add(assistant_msg)
                                session2.commit()
                            finally:
                                session2.close()
                        except Exception as db_err:
                            logger.error(f"Async DB save error: {db_err}")

                    threading.Thread(target=save_to_db, daemon=True).start()

            except GeneratorExit:
                logger.info("Stream cancelled by client")
                return
            except Exception as e:
                logger.error(f"Streaming error: {e}")
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
                yield "data: [DONE]\n\n"

        return Response(stream_with_context(generate()), mimetype='text/event-stream',
            headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive', 'X-Accel-Buffering': 'no'})

    def get_conversations(self):
        session = self.db.get_session()
        conversations = session.query(self.Conversation).order_by(self.Conversation.updated_at.desc()).limit(50).all()
        result = [{'id': conv.id, 'title': conv.title, 'created_at': conv.created_at.isoformat(), 'updated_at': conv.updated_at.isoformat()} for conv in conversations]
        session.close()
        return jsonify({'success': True, 'data': result})

    def get_conversation(self, conversation_id):
        session = self.db.get_session()
        messages = session.query(self.Message).filter_by(conversation_id=conversation_id).order_by(self.Message.created_at).all()
        result = [{'id': msg.id, 'role': msg.role, 'content': msg.content, 'citations': msg.citations, 'created_at': msg.created_at.isoformat()} for msg in messages]
        session.close()
        return jsonify({'success': True, 'data': result})

    def delete_conversation(self, conversation_id):
        session = self.db.get_session()
        session.query(self.Message).filter_by(conversation_id=conversation_id).delete()
        session.query(self.Conversation).filter_by(id=conversation_id).delete()
        session.commit()
        session.close()
        return jsonify({'success': True})

    def get_documents(self):
        session = self.db.get_session()
        status_filter = request.args.get('status')
        query = session.query(self.Document)
        if status_filter:
            query = query.filter_by(status=status_filter)
        documents = query.all()
        result = [{'id': doc.id, 'name': doc.name, 'source': doc.source, 'status': doc.status,
            'chunk_count': doc.chunk_count, 'file_size': doc.file_size,
            'upload_time': doc.created_at.isoformat(), 'created_at': doc.created_at.isoformat()} for doc in documents]
        session.close()
        return jsonify({'success': True, 'data': result})

    def search_documents(self):
        keyword = request.args.get('keyword', '').strip()
        if not keyword:
            return jsonify({'success': True, 'data': []})
        session = self.db.get_session()
        documents = session.query(self.Document).filter(self.Document.name.contains(keyword)).all()
        result = [{'id': doc.id, 'name': doc.name, 'source': doc.source, 'status': doc.status,
            'chunk_count': doc.chunk_count, 'file_size': doc.file_size,
            'upload_time': doc.created_at.isoformat(), 'created_at': doc.created_at.isoformat()} for doc in documents]
        session.close()
        return jsonify({'success': True, 'data': result})

    def upload_document(self):
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': {'code': 'NO_FILE', 'message': '未找到上传文件'}}), 400
        file = request.files['file']
        if not file.filename.endswith(('.pdf', '.docx', '.txt', '.rst')):
            return jsonify({'success': False, 'error': {'code': 'INVALID_FILE_TYPE', 'message': '不支持的文件格式'}}), 400

        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        if file_size > MAX_FILE_SIZE:
            return jsonify({'success': False, 'error': {'code': 'FILE_TOO_LARGE', 'message': '文件大小超过限制（最大50MB）'}}), 400
        if file_size == 0:
            return jsonify({'success': False, 'error': {'code': 'EMPTY_FILE', 'message': '文件为空'}}), 400

        filename = f"{uuid.uuid4()}_{file.filename}"
        filepath = os.path.join(Config.UPLOAD_PATH, filename)
        os.makedirs(Config.UPLOAD_PATH, exist_ok=True)
        with open(filepath, 'wb') as f:
            chunk_size = 1024 * 1024
            while True:
                chunk = file.read(chunk_size)
                if not chunk:
                    break
                f.write(chunk)

        doc_id = str(uuid.uuid4())
        session = self.db.get_session()
        document = self.Document(id=doc_id, name=file.filename, source='uploaded', file_path=filepath,
            status='uploading', file_size=file_size, upload_progress=0, uploaded_bytes=file_size, total_bytes=file_size)
        session.add(document)
        session.commit()
        session.close()

        def process_with_timeout():
            thread = threading.Thread(target=self._process_document_async, args=(doc_id, filepath, file.filename))
            thread.daemon = True
            thread.start()
            thread.join(timeout=120)
            if thread.is_alive():
                try:
                    session = self.db.get_session()
                    doc = session.query(self.Document).filter_by(id=doc_id).first()
                    if doc and doc.status == 'processing':
                        doc.status = 'failed'
                        doc.error_message = '处理超时（超过120秒）'
                        session.commit()
                    session.close()
                except:
                    pass
                upload_tasks[doc_id] = {'progress': 0, 'status': 'failed', 'error': '处理超时'}

        threading.Thread(target=process_with_timeout, daemon=True).start()
        return jsonify({'success': True, 'data': {'id': doc_id, 'name': file.filename, 'status': 'processing', 'message': '文档上传成功，正在后台处理中...'}})

    def _process_document_async(self, doc_id, filepath, filename):
        try:
            session = self.db.get_session()
            doc = session.query(self.Document).filter_by(id=doc_id).first()
            if doc:
                doc.status = 'processing'
                doc.upload_progress = 10
                session.commit()
            session.close()
            upload_tasks[doc_id] = {'progress': 10, 'status': 'processing'}

            processor = DocumentProcessor()
            if filename.endswith('.pdf'):
                upload_tasks[doc_id] = {'progress': 30, 'status': 'processing'}
                chunks = processor.process_pdf(filepath)
            elif filename.endswith('.docx'):
                upload_tasks[doc_id] = {'progress': 30, 'status': 'processing'}
                chunks = processor.process_docx(filepath)
            elif filename.endswith('.rst'):
                upload_tasks[doc_id] = {'progress': 30, 'status': 'processing'}
                chunks = processor.process_rst_file(filepath)
            else:
                upload_tasks[doc_id] = {'progress': 30, 'status': 'processing'}
                chunks = processor.process_txt(filepath)

            upload_tasks[doc_id] = {'progress': 60, 'status': 'processing'}
            if chunks:
                self.vector_store.add_chunks(chunks)
            upload_tasks[doc_id] = {'progress': 90, 'status': 'processing'}

            session = self.db.get_session()
            doc = session.query(self.Document).filter_by(id=doc_id).first()
            if doc:
                doc.status = 'completed'
                doc.chunk_count = len(chunks) if chunks else 0
                doc.upload_progress = 100
            session.commit()
            session.close()
            upload_tasks[doc_id] = {'progress': 100, 'status': 'completed'}
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            try:
                session = self.db.get_session()
                doc = session.query(self.Document).filter_by(id=doc_id).first()
                if doc:
                    doc.status = 'failed'
                    doc.error_message = str(e)
                session.commit()
                session.close()
            except:
                pass
            upload_tasks[doc_id] = {'progress': 0, 'status': 'failed', 'error': str(e)}

    def get_upload_progress(self, doc_id):
        session = self.db.get_session()
        doc = session.query(self.Document).filter_by(id=doc_id).first()
        if not doc:
            session.close()
            return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '文档不存在'}}), 404
        result = {'id': doc.id, 'name': doc.name, 'status': doc.status, 'progress': doc.upload_progress,
            'uploaded_bytes': doc.uploaded_bytes, 'total_bytes': doc.total_bytes,
            'chunk_count': doc.chunk_count, 'error_message': doc.error_message}
        session.close()
        return jsonify({'success': True, 'data': result})

    def delete_document(self, doc_id):
        session = self.db.get_session()
        doc = session.query(self.Document).filter_by(id=doc_id).first()
        if doc:
            if doc.file_path:
                self.vector_store.delete_by_source(doc.file_path)
            if os.path.exists(doc.file_path):
                os.remove(doc.file_path)
            session.delete(doc)
            session.commit()
        session.close()
        return jsonify({'success': True})

    def get_document_preview(self, doc_id):
        session = self.db.get_session()
        doc = session.query(self.Document).filter_by(id=doc_id).first()
        if not doc:
            session.close()
            return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '文档不存在'}}), 404

        preview_text = ""
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                ext = os.path.splitext(doc.file_path)[1].lower()
                with open(doc.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    if ext in ('.txt', '.rst'):
                        preview_text = f.read()[:5000]
                    elif ext == '.pdf':
                        f.close()
                        try:
                            from pypdf import PdfReader
                            reader = PdfReader(doc.file_path)
                            text_parts = [p.extract_text() for p in reader.pages[:10] if p.extract_text()]
                            preview_text = '\n'.join(text_parts)[:5000]
                        except Exception as e:
                            preview_text = f"[PDF预览失败: {str(e)}]"
                    elif ext == '.docx':
                        f.close()
                        try:
                            from docx import Document as DocxDocument
                            docx_doc = DocxDocument(doc.file_path)
                            paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
                            preview_text = '\n'.join(paragraphs)[:5000]
                        except Exception as e:
                            preview_text = f"[DOCX预览失败: {str(e)}]"
                    else:
                        preview_text = "[不支持的文件格式]"
            except Exception as e:
                preview_text = f"[读取失败: {str(e)}]"
        else:
            preview_text = "[文件不存在或路径为空]"

        file_size = os.path.getsize(doc.file_path) if doc.file_path and os.path.exists(doc.file_path) else 0
        session.close()
        return jsonify({'success': True, 'data': {
            'id': doc.id, 'name': doc.name, 'source': doc.source, 'status': doc.status,
            'chunk_count': doc.chunk_count, 'file_size': file_size,
            'preview_text': preview_text, 'created_at': doc.created_at.isoformat()
        }})

    def get_document_chunks(self, doc_id):
        session = self.db.get_session()
        doc = session.query(self.Document).filter_by(id=doc_id).first()
        if not doc:
            session.close()
            return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '文档不存在'}}), 404

        chunks = []
        if self.vector_store and doc.file_path:
            try:
                results = self.vector_store.collection.query(
                    query_embeddings=[[0.0] * self.vector_store.embedding_dimension],
                    n_results=doc.chunk_count or 10, where={"source": doc.file_path})
                if results and 'documents' in results:
                    for i, chunk_text in enumerate(results.get('documents', [[]])[0]):
                        chunk_id = results.get('ids', [[]])[0][i] if results.get('ids') else f"chunk_{i}"
                        metadata = results.get('metadatas', [[{}]])[0][i] if results.get('metadatas') else {}
                        chunks.append({'id': chunk_id, 'text': chunk_text, 'chunk_index': metadata.get('chunk_index', i)})
            except Exception as e:
                logger.error(f"Get chunks failed: {e}")
        session.close()
        return jsonify({'success': True, 'data': chunks})

    def submit_feedback(self):
        data = request.get_json(silent=True) or {}
        message_id = data.get('message_id')
        rating = data.get('rating')
        conversation_id = data.get('conversation_id')
        reason = data.get('reason')
        reason_category = data.get('reason_category', '')
        custom_reason_text = data.get('custom_reason_text', '')

        if not message_id:
            return jsonify({'success': False, 'error': {'code': 'INVALID_REQUEST', 'message': 'message_id 不能为空'}}), 400
        if rating not in ('positive', 'negative'):
            return jsonify({'success': False, 'error': {'code': 'INVALID_REQUEST', 'message': 'rating 必须是 positive 或 negative'}}), 400

        if rating == 'negative' and reason_category:
            if reason_category not in FEEDBACK_REASON_CATEGORIES:
                return jsonify({'success': False, 'error': {'code': 'INVALID_REASON', 'message': f'无效的原因分类，必须是: {", ".join(FEEDBACK_REASON_CATEGORIES)}'}}), 400

        session = self.db.get_session()
        try:
            existing = session.query(self.Feedback).filter_by(message_id=message_id).first()
            if existing:
                return jsonify({'success': False, 'error': {'code': 'DUPLICATE_FEEDBACK', 'message': '该消息已被评价'}}), 409

            stored_reason = reason_category if reason_category else reason
            if reason_category == 'other' and custom_reason_text:
                stored_reason = f"other: {custom_reason_text}"

            feedback = self.Feedback(
                id=str(uuid.uuid4()),
                message_id=message_id,
                conversation_id=conversation_id or '',
                rating=rating,
                reason=stored_reason,
                game_type=self.game_type,
                feedback_text=data.get('feedback_text'),
            )
            session.add(feedback)

            if rating == 'negative':
                alert = self.Alert(id=str(uuid.uuid4()), rule_type='negative_feedback', message=f'收到差评: message_id={message_id}, reason={stored_reason or "无"}')
                session.add(alert)

                logger.info(f"[Feedback] Negative feedback received: message_id={message_id}, reason={stored_reason}")

                try:
                    user_msg = session.query(self.Message).filter_by(conversation_id=conversation_id, role='user').order_by(self.Message.created_at.desc()).first()
                    assistant_msg = session.query(self.Message).filter_by(id=message_id, role='assistant').first()

                    if user_msg and assistant_msg and self.NegativeSample:
                        negative_sample = self.NegativeSample(
                            id=str(uuid.uuid4()),
                            question=user_msg.content,
                            answer=assistant_msg.content,
                            reason=stored_reason or 'unspecified',
                            feedback_id=feedback.id
                        )
                        session.add(negative_sample)
                        logger.info(f"[NegativeSample] Auto-added to negative sample库: feedback_id={feedback.id}")
                except Exception as e:
                    logger.error(f"[NegativeSample] Failed to add negative sample: {e}")

            session.commit()
            return jsonify({'success': True, 'data': {'id': feedback.id}})
        finally:
            session.close()

    def get_feedback_status(self, message_id):
        session = self.db.get_session()
        try:
            feedback = session.query(self.Feedback).filter_by(message_id=message_id).first()
            if feedback:
                return jsonify({'success': True, 'data': {'id': feedback.id, 'message_id': feedback.message_id, 'rating': feedback.rating, 'reason': feedback.reason, 'created_at': feedback.created_at.isoformat()}})
            return jsonify({'success': True, 'data': {'exists': False}})
        finally:
            session.close()

    def get_feedback_stats(self):
        session = self.db.get_session()
        try:
            total = session.query(self.Feedback).filter_by(game_type=self.game_type).count()
            positive = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='positive').count()
            negative = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').count()
            rate = round(positive / total, 4) if total > 0 else 0.0

            negative_feedbacks = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').filter(self.Feedback.reason.isnot(None)).filter(self.Feedback.reason != '').all()
            reason_counts = Counter([f.reason for f in negative_feedbacks])
            reason_distribution = [{'reason': r, 'count': c} for r, c in reason_counts.most_common(10)]

            negative_sample_count = 0
            if self.NegativeSample:
                negative_sample_count = session.query(self.NegativeSample).count()

            return jsonify({'success': True, 'data': {
                'total_feedbacks': total,
                'positive_count': positive,
                'negative_count': negative,
                'positive_rate': rate,
                'reason_distribution': reason_distribution,
                'negative_sample_count': negative_sample_count
            }})
        finally:
            session.close()

    def get_negative_samples(self):
        if not self.NegativeSample:
            return jsonify({'success': False, 'error': {'code': 'NOT_AVAILABLE', 'message': '负样本库未初始化'}}), 503

        session = self.db.get_session()
        try:
            page = request.args.get('page', 1, type=int)
            limit = request.args.get('limit', 20, type=int)
            reason_filter = request.args.get('reason', '')

            query = session.query(self.NegativeSample)
            if reason_filter:
                query = query.filter(self.NegativeSample.reason == reason_filter)

            total = query.count()
            samples = query.order_by(self.NegativeSample.created_at.desc()).offset((page - 1) * limit).limit(limit).all()

            result = [{
                'id': s.id,
                'question': s.question,
                'answer': s.answer,
                'reason': s.reason,
                'feedback_id': s.feedback_id,
                'created_at': s.created_at.isoformat()
            } for s in samples]

            return jsonify({'success': True, 'data': {
                'samples': result,
                'total': total,
                'page': page,
                'limit': limit
            }})
        finally:
            session.close()

    def _check_alerts(self, session, total_conversations, total_messages, knowledge_base_size):
        rules = session.query(self.AlertRule).filter_by(enabled=1).all()
        rule_map = {r.rule_type: r for r in rules}
        default_rules = [
            {'rule_type': 'low_knowledge_base', 'threshold': 100, 'description': '知识库块数低于阈值'},
            {'rule_type': 'no_conversations', 'threshold': 1, 'description': '无对话记录'},
        ]
        for rule_def in default_rules:
            rule_type = rule_def['rule_type']
            if rule_type not in rule_map:
                new_rule = self.AlertRule(id=str(uuid.uuid4()), rule_type=rule_type, threshold=rule_def['threshold'], enabled=1, description=rule_def['description'])
                session.add(new_rule)
                session.commit()
                rule_map[rule_type] = new_rule

        if 'low_knowledge_base' in rule_map:
            rule = rule_map['low_knowledge_base']
            if knowledge_base_size < rule.threshold:
                existing = session.query(self.Alert).filter_by(rule_type='low_knowledge_base', is_read=0).first()
                if not existing:
                    alert = self.Alert(id=str(uuid.uuid4()), rule_type='low_knowledge_base', message=f'知识库块数 ({knowledge_base_size}) 低于阈值 ({rule.threshold})')
                    session.add(alert)
                    session.commit()
                    self._send_feishu_notification('知识库告警', f'知识库块数 ({knowledge_base_size}) 低于阈值 ({rule.threshold})', 'high')

        if 'no_conversations' in rule_map:
            rule = rule_map['no_conversations']
            if total_conversations < rule.threshold:
                existing = session.query(self.Alert).filter_by(rule_type='no_conversations', is_read=0).first()
                if not existing:
                    alert = self.Alert(id=str(uuid.uuid4()), rule_type='no_conversations', message=f'暂无对话记录 ({total_conversations}/{rule.threshold})')
                    session.add(alert)
                    session.commit()
                    self._send_feishu_notification('系统告警', f'暂无对话记录 ({total_conversations}/{rule.threshold})', 'medium')

    def _send_feishu_notification(self, title, content, priority='medium'):
        try:
            if not Config.FEISHU_ENABLED:
                return
            if not check_feishu_cli_available():
                return
            result = feishu_send_alert(title, content, priority)
            if result.get('success'):
                logger.info(f"Feishu alert sent: {title}")
        except Exception as e:
            logger.error(f"Feishu alert error: {e}")

    def get_metrics(self):
        session = self.db.get_session()
        total_conversations = session.query(self.Conversation).count()
        total_messages = session.query(self.Message).count()
        stats = self.vector_store.get_collection_stats() if self.vector_store else {}
        knowledge_base_size = stats.get('count', 0)
        self._check_alerts(session, total_conversations, total_messages, knowledge_base_size)
        unread_count = session.query(self.Alert).filter_by(is_read=0).count()
        today_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        messages_today = session.query(self.Message).filter(self.Message.created_at >= today_start).count()
        week_start = today_start - timedelta(days=today_start.weekday())
        conversations_this_week = session.query(self.Conversation).filter(self.Conversation.created_at >= week_start).count()
        logs = session.query(self.PerformanceLog).filter(self.PerformanceLog.created_at >= today_start).all()
        if logs:
            avg_response_time_ms = int(sum(log.response_time_ms for log in logs) / len(logs))
        else:
            avg_response_time_ms = 1500 if session.query(self.Message).filter_by(role='assistant').first() else 0

        user_messages = session.query(self.Message).filter_by(role='user').all()
        top_keywords = []
        if user_messages:
            stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这', '吗', '什么', '怎么', '为', '么', '呢', '吧', '啊', '哪', '谁', '如何', '怎样', '为什么', '能不能', '可不可以', '请', '问', '想', '知道'}
            all_words = []
            for msg in user_messages:
                words = re.findall(r'[\u4e00-\u9fa5]{2,}', msg.content)
                words = [w for w in words if w not in stop_words and len(w) >= 2]
                all_words.extend(words)
            word_counts = Counter(all_words)
            top_keywords = [{'keyword': word, 'count': count} for word, count in word_counts.most_common(10)]
        session.close()

        return jsonify({'success': True, 'data': {
            'total_conversations': total_conversations, 'total_messages': total_messages,
            'knowledge_base_size': knowledge_base_size, 'unread_alerts': unread_count,
            'avg_response_time_ms': avg_response_time_ms, 'messages_today': messages_today,
            'conversations_this_week': conversations_this_week, 'top_keywords': top_keywords
        }})

    def get_alert_rules(self):
        session = self.db.get_session()
        rules = session.query(self.AlertRule).all()
        result = [{'id': r.id, 'rule_type': r.rule_type, 'threshold': r.threshold, 'enabled': r.enabled, 'description': r.description} for r in rules]
        session.close()
        return jsonify({'success': True, 'data': result})

    def create_or_update_alert_rule(self):
        data = request.get_json(silent=True) or {}
        rule_type = data.get('rule_type')
        if not rule_type:
            return jsonify({'success': False, 'error': {'code': 'INVALID_RULE', 'message': '规则类型不能为空'}}), 400
        session = self.db.get_session()
        rule = session.query(self.AlertRule).filter_by(rule_type=rule_type).first()
        if rule:
            rule.threshold = data.get('threshold', 0)
            rule.enabled = data.get('enabled', 1)
            rule.description = data.get('description', '')
        else:
            rule = self.AlertRule(id=str(uuid.uuid4()), rule_type=rule_type, threshold=data.get('threshold', 0), enabled=data.get('enabled', 1), description=data.get('description', ''))
            session.add(rule)
        session.commit()
        session.close()
        return jsonify({'success': True, 'data': {'rule_type': rule_type}})

    def get_alert_history(self):
        session = self.db.get_session()
        limit = request.args.get('limit', 50, type=int)
        alerts = session.query(self.Alert).order_by(self.Alert.created_at.desc()).limit(limit).all()
        result = [{'id': a.id, 'rule_type': a.rule_type, 'message': a.message, 'is_read': a.is_read, 'created_at': a.created_at.isoformat()} for a in alerts]
        session.close()
        return jsonify({'success': True, 'data': result})

    def mark_alert_read(self, alert_id):
        session = self.db.get_session()
        alert = session.query(self.Alert).filter_by(id=alert_id).first()
        if not alert:
            session.close()
            return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '告警不存在'}}), 404
        alert.is_read = 1
        session.commit()
        session.close()
        return jsonify({'success': True})

    def get_unread_alerts(self):
        session = self.db.get_session()
        alerts = session.query(self.Alert).filter_by(is_read=0).order_by(self.Alert.created_at.desc()).all()
        result = [{'id': a.id, 'rule_type': a.rule_type, 'message': a.message, 'created_at': a.created_at.isoformat()} for a in alerts]
        session.close()
        return jsonify({'success': True, 'data': result})

    def get_metrics_trend(self):
        days = min(request.args.get('days', 7, type=int), 30)
        session = self.db.get_session()
        end_date = datetime.utcnow().replace(hour=23, minute=59, second=59, microsecond=999999)
        start_date = (end_date - timedelta(days=days - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
        daily_data = []
        current_date = start_date
        while current_date <= end_date:
            day_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = current_date.replace(hour=23, minute=59, second=59, microsecond=999999)
            mc = session.query(self.Message).filter(self.Message.created_at >= day_start, self.Message.created_at <= day_end).count()
            cc = session.query(self.Conversation).filter(self.Conversation.created_at >= day_start, self.Conversation.created_at <= day_end).count()
            daily_data.append({'date': day_start.strftime('%Y-%m-%d'), 'messages': mc, 'conversations': cc})
            current_date += timedelta(days=1)
        session.close()
        return jsonify({'success': True, 'data': {'days': days, 'data': daily_data}})

    def get_quality_metrics(self):
        session = self.db.get_session()
        try:
            total = session.query(self.Feedback).filter_by(game_type=self.game_type).count()
            positive = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='positive').count()
            negative = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').count()
            rate = positive / total if total > 0 else 0.0
            end_date = datetime.utcnow().replace(hour=23, minute=59, second=59, microsecond=999999)
            start_date = (end_date - timedelta(days=29)).replace(hour=0, minute=0, second=0, microsecond=0)
            daily_trend = []
            current_date = start_date
            while current_date <= end_date:
                day_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0)
                day_end = current_date.replace(hour=23, minute=59, second=59, microsecond=999999)
                dp = session.query(self.Feedback).filter(self.Feedback.game_type == self.game_type, self.Feedback.rating == 'positive', self.Feedback.created_at >= day_start, self.Feedback.created_at <= day_end).count()
                dn = session.query(self.Feedback).filter(self.Feedback.game_type == self.game_type, self.Feedback.rating == 'negative', self.Feedback.created_at >= day_start, self.Feedback.created_at <= day_end).count()
                daily_trend.append({'date': day_start.strftime('%Y-%m-%d'), 'positive': dp, 'negative': dn})
                current_date += timedelta(days=1)
            neg_feedbacks = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').filter(self.Feedback.reason.isnot(None)).filter(self.Feedback.reason != '').all()
            reason_counts = Counter([f.reason for f in neg_feedbacks])
            top_reasons = [{'reason': r, 'count': c} for r, c in reason_counts.most_common(10)]
            return jsonify({'success': True, 'data': {
                'total_feedbacks': total, 'positive_count': positive, 'negative_count': negative,
                'positive_rate': round(rate, 4), 'daily_trend': daily_trend, 'top_negative_reasons': top_reasons
            }})
        finally:
            session.close()

    def get_search_quality(self):
        days = min(request.args.get('days', 7, type=int), 30)
        session = self.db.get_session()
        try:
            end_date = datetime.utcnow().replace(hour=23, minute=59, second=59, microsecond=999999)
            start_date = (end_date - timedelta(days=days - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
            total = session.query(self.Feedback).filter_by(game_type=self.game_type).count()
            positive = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='positive').count()
            negative = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').count()
            rate = positive / total if total > 0 else 0.0
            daily_quality = []
            current_date = start_date
            while current_date <= end_date:
                day_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0)
                day_end = current_date.replace(hour=23, minute=59, second=59, microsecond=999999)
                dp = session.query(self.Feedback).filter(self.Feedback.game_type == self.game_type, self.Feedback.rating == 'positive', self.Feedback.created_at >= day_start, self.Feedback.created_at <= day_end).count()
                dn = session.query(self.Feedback).filter(self.Feedback.game_type == self.game_type, self.Feedback.rating == 'negative', self.Feedback.created_at >= day_start, self.Feedback.created_at <= day_end).count()
                dt = dp + dn
                daily_quality.append({'date': day_start.strftime('%Y-%m-%d'), 'positive': dp, 'negative': dn, 'rate': round(dp / dt, 2) if dt > 0 else None})
                current_date += timedelta(days=1)
            neg_feedbacks = session.query(self.Feedback).filter_by(game_type=self.game_type, rating='negative').filter(self.Feedback.reason.isnot(None)).filter(self.Feedback.reason != '').all()
            reason_counts = Counter([f.reason for f in neg_feedbacks])
            top_reasons = [{'reason': r, 'count': c} for r, c in reason_counts.most_common(10)]
            model_logs = session.query(self.PerformanceLog).filter(self.PerformanceLog.created_at >= start_date).order_by(self.PerformanceLog.response_time_ms).all()
            if model_logs:
                all_times = [log.response_time_ms for log in model_logs]
                avg_rt = int(sum(all_times) / len(all_times))
                dist = {'fast': sum(1 for t in all_times if t < 1000), 'medium': sum(1 for t in all_times if 1000 <= t <= 3000), 'slow': sum(1 for t in all_times if t > 3000)}
            else:
                avg_rt, dist = 0, {'fast': 0, 'medium': 0, 'slow': 0}
            return jsonify({'success': True, 'data': {
                'days': days, 'positive_rate': round(rate, 4), 'total_feedbacks': total,
                'positive_count': positive, 'negative_count': negative, 'daily_quality': daily_quality,
                'top_negative_reasons': top_reasons, 'avg_response_time_ms': avg_rt, 'response_distribution': dist
            }})
        finally:
            session.close()

    def get_cache_stats(self):
        """获取缓存层统计信息（L1 + L2）"""
        if self.rag_engine and hasattr(self.rag_engine, 'get_cache_stats'):
            cache_stats = self.rag_engine.get_cache_stats()
        else:
            from app.services.query_cache import get_l1_cache
            from app.services.vector_cache import get_query_cache
            cache_stats = {
                'l1_cache': get_l1_cache().get_stats(),
                'l2_cache': get_query_cache().get_stats()
            }
        return jsonify({'success': True, 'data': cache_stats})

    def get_rag_config(self):
        template_names = {'default': f'默认 ({self.game_type.upper()})', 'concise': '简洁模式', 'detailed': '详细模式'}
        templates_list = [{'id': k, 'name': template_names.get(k, k)} for k in self.prompt_templates.keys()]
        return jsonify({'success': True, 'data': {'config': dict(self.rag_runtime_config), 'available_templates': templates_list}})

    def update_rag_config(self):
        data = request.get_json(silent=True) or {}
        if not data:
            return jsonify({'success': False, 'error': {'code': 'INVALID_REQUEST', 'message': '请求体不能为空'}}), 400
        allowed_keys = {'top_k', 'temperature', 'max_tokens', 'system_prompt_template', 'streaming_enabled', 'similarity_threshold'}
        for key in data:
            if key not in allowed_keys:
                return jsonify({'success': False, 'error': {'code': 'INVALID_KEY', 'message': f'无效的配置项: {key}'}}), 400
        if 'top_k' in data:
            if not isinstance(data['top_k'], int) or data['top_k'] < 1 or data['top_k'] > 20:
                return jsonify({'success': False, 'error': {'code': 'INVALID_VALUE', 'message': 'top_k 必须是 1-20 之间的整数'}}), 400
            self.rag_runtime_config['top_k'] = data['top_k']
        if 'temperature' in data:
            if not isinstance(data['temperature'], (int, float)) or data['temperature'] < 0.0 or data['temperature'] > 1.0:
                return jsonify({'success': False, 'error': {'code': 'INVALID_VALUE', 'message': 'temperature 必须是 0.0-1.0 之间的数值'}}), 400
            self.rag_runtime_config['temperature'] = float(data['temperature'])
        if 'max_tokens' in data:
            if not isinstance(data['max_tokens'], int) or data['max_tokens'] < 100 or data['max_tokens'] > 4000:
                return jsonify({'success': False, 'error': {'code': 'INVALID_VALUE', 'message': 'max_tokens 必须是 100-4000 之间的整数'}}), 400
            self.rag_runtime_config['max_tokens'] = data['max_tokens']
        if 'system_prompt_template' in data:
            if data['system_prompt_template'] not in self.prompt_templates:
                return jsonify({'success': False, 'error': {'code': 'INVALID_VALUE', 'message': f'无效的模板: {data["system_prompt_template"]}'}}), 400
            self.rag_runtime_config['system_prompt_template'] = data['system_prompt_template']
        if 'streaming_enabled' in data:
            self.rag_runtime_config['streaming_enabled'] = bool(data['streaming_enabled'])
        if 'similarity_threshold' in data:
            if not isinstance(data['similarity_threshold'], (int, float)) or data['similarity_threshold'] < 0.0 or data['similarity_threshold'] > 1.0:
                return jsonify({'success': False, 'error': {'code': 'INVALID_VALUE', 'message': 'similarity_threshold 必须是 0.0-1.0 之间的数值'}}), 400
            self.rag_runtime_config['similarity_threshold'] = float(data['similarity_threshold'])
        return jsonify({'success': True, 'data': {'config': dict(self.rag_runtime_config)}})

    def agent_query(self):
        data = request.get_json(silent=True) or {}
        user_input = data.get('question', '')
        if not user_input:
            return jsonify({'success': False, 'error': {'code': 'EMPTY_QUESTION'}}), 400
        if self.agent_instance is None:
            return jsonify({'success': False, 'error': {'code': 'AGENT_NOT_INITIALIZED'}}), 503
        result = self.agent_instance.run(user_input)

        agent_type = type(self.agent_instance).__name__
        is_function_calling = agent_type == 'FunctionCallingAgent'

        # 兼容两种 Agent 的返回格式
        thoughts = result.get('thoughts', [])
        actions = result.get('actions', [])
        tool_calls_history = result.get('tool_calls_history', [])

        # 派生 actions 列表（即使 ReAct 没填，也保证前端能看到调用链）
        if not actions and tool_calls_history:
            actions = [
                tc.get('function', {}).get('name', '')
                if isinstance(tc, dict) else getattr(tc, 'function', {}).get('name', '')
                for tc in tool_calls_history
            ]
        if not thoughts and tool_calls_history:
            tool_names = [
                tc.get('function', {}).get('name', '')
                if isinstance(tc, dict) else getattr(tc, 'function', {}).get('name', '')
                for tc in tool_calls_history
            ]
            thoughts = [f"调用了工具: {', '.join(tool_names)}"]

        return jsonify({
            'success': True,
            'data': {
                'answer': result.get('answer', ''),
                'thoughts': thoughts,
                'actions': actions,
                'tool_calls_history': tool_calls_history,
                'agent_type': agent_type,
                'used_function_calling': is_function_calling,
            }
        })

    def list_skills(self):
        if self.skill_registry is None:
            return jsonify({'success': False, 'error': {'code': 'AGENT_NOT_INITIALIZED'}}), 503
        skills = self.skill_registry.get_available_skills()
        return jsonify({'success': True, 'data': {'skills': skills}})

    def multimodal_query(self):
        data = request.get_json(silent=True) or {}
        text = data.get('question', '')
        image_base64 = data.get('image_base64', None)
        top_k = data.get('top_k', 5)
        if not text:
            return jsonify({'success': False, 'error': {'code': 'EMPTY_QUESTION'}}), 400
        if self.rag_engine is None:
            return jsonify({'success': False, 'error': {'code': 'RAG_NOT_INITIALIZED'}}), 503
        response = self.rag_engine.query_multimodal(text, image_base64, top_k)
        return jsonify({'success': True, 'data': {'answer': response.answer, 'citations': response.citations, 'confidence': response.confidence, 'has_vision_analysis': image_base64 is not None}})

    def import_url(self):
        """网页URL抓取并入库"""
        data = request.get_json(silent=True) or {}
        url = data.get('url', '').strip()
        title = data.get('title', '')
        selectors = data.get('selectors')
        auto_chunk = data.get('auto_chunk', True)

        if not url:
            return jsonify({'success': False, 'error': {'code': 'INVALID_URL', 'message': 'URL不能为空'}}), 400

        import re
        if not re.match(r'^https?://', url):
            return jsonify({'success': False, 'error': {'code': 'INVALID_URL', 'message': 'URL格式不正确，必须以http://或https://开头'}}), 400

        doc_id = str(uuid.uuid4())
        session = self.db.get_session()
        try:
            document = self.Document(
                id=doc_id,
                name=title or url.split('/')[-1] or '网页导入',
                source='url_import',
                file_path=url,
                status='processing',
                chunk_count=0,
                file_size=0,
                upload_progress=10,
                uploaded_bytes=0,
                total_bytes=0
            )
            session.add(document)
            session.commit()
        finally:
            session.close()

        def process_url_async():
            try:
                import asyncio
                processor = DocumentProcessor()
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                fetch_result = loop.run_until_complete(processor.fetch_url_content(url, selectors))
                loop.close()

                session = self.db.get_session()
                doc = session.query(self.Document).filter_by(id=doc_id).first()
                if not fetch_result.get('success'):
                    if doc:
                        doc.status = 'failed'
                        doc.error_message = fetch_result.get('error', '抓取失败')
                        doc.upload_progress = 0
                    session.commit()
                    session.close()
                    return

                if doc:
                    doc.upload_progress = 50
                session.commit()
                session.close()

                content = fetch_result.get('content', '')
                fetched_title = fetch_result.get('title', title or url)

                if auto_chunk and len(content) >= 50:
                    chunks = processor.process_url_content(content, fetched_title, url)
                    if chunks:
                        self.vector_store.add_chunks(chunks)

                session = self.db.get_session()
                doc = session.query(self.Document).filter_by(id=doc_id).first()
                if doc:
                    doc.status = 'completed'
                    doc.chunk_count = len(chunks) if chunks else 0
                    doc.name = fetched_title
                    doc.file_size = len(content)
                    doc.upload_progress = 100
                session.commit()
                session.close()
            except Exception as e:
                logger.error(f"URL import failed: {e}")
                try:
                    session = self.db.get_session()
                    doc = session.query(self.Document).filter_by(id=doc_id).first()
                    if doc:
                        doc.status = 'failed'
                        doc.error_message = str(e)
                    session.commit()
                    session.close()
                except:
                    pass

        threading.Thread(target=process_url_async, daemon=True).start()
        return jsonify({'success': True, 'data': {
            'id': doc_id,
            'url': url,
            'status': 'processing',
            'message': '网页抓取任务已提交，正在后台处理中...'
        }})

    def get_import_status(self, doc_id):
        """获取URL导入状态"""
        session = self.db.get_session()
        try:
            doc = session.query(self.Document).filter_by(id=doc_id).first()
            if not doc:
                return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '任务不存在'}}), 404
            return jsonify({'success': True, 'data': {
                'id': doc.id,
                'name': doc.name,
                'status': doc.status,
                'progress': doc.upload_progress,
                'chunk_count': doc.chunk_count,
                'error_message': doc.error_message,
                'source_url': doc.file_path if doc.source == 'url_import' else None
            }})
        finally:
            session.close()

    def search_conversations(self):
        """对话历史搜索"""
        keyword = request.args.get('q', '').strip()
        page = request.args.get('page', 1, type=int)
        limit = request.args.get('limit', 20, type=int)

        if not keyword:
            return jsonify({'success': True, 'data': {'results': [], 'total': 0, 'page': 1}})

        session = self.db.get_session()
        try:
            messages = session.query(self.Message).filter(
                self.Message.content.contains(keyword)
            ).order_by(self.Message.created_at.desc()).all()

            conv_ids = list(set([msg.conversation_id for msg in messages]))
            conversations = session.query(self.Conversation).filter(
                self.Conversation.id.in_(conv_ids)
            ).all()
            conv_map = {c.id: c for c in conversations}

            results = []
            for msg in messages:
                conv = conv_map.get(msg.conversation_id)
                if conv:
                    matched_text = msg.content
                    pos = matched_text.find(keyword)
                    if pos != -1:
                        start = max(0, pos - 50)
                        end = min(len(matched_text), pos + len(keyword) + 50)
                        context = matched_text[start:end]
                        if start > 0:
                            context = '...' + context
                        if end < len(matched_text):
                            context = context + '...'
                    else:
                        context = matched_text[:100]

                    results.append({
                        'conversation_id': msg.conversation_id,
                        'conversation_title': conv.title,
                        'message_id': msg.id,
                        'role': msg.role,
                        'matched_context': context,
                        'created_at': msg.created_at.isoformat()
                    })

            total = len(results)
            start_idx = (page - 1) * limit
            end_idx = start_idx + limit
            paginated_results = results[start_idx:end_idx]

            return jsonify({'success': True, 'data': {
                'results': paginated_results,
                'total': total,
                'page': page,
                'limit': limit
            }})
        finally:
            session.close()

    def get_citation_detail(self, chunk_id):
        """获取引用原文详情"""
        if not self.vector_store:
            return jsonify({'success': False, 'error': {'code': 'NOT_AVAILABLE', 'message': '向量库未初始化'}}), 503

        try:
            results = self.vector_store.collection.get(ids=[chunk_id])
            if not results or not results.get('documents'):
                return jsonify({'success': False, 'error': {'code': 'NOT_FOUND', 'message': '引用内容不存在'}}), 404

            doc_text = results['documents'][0]
            metadata = results.get('metadatas', [{}])[0]

            return jsonify({'success': True, 'data': {
                'chunk_id': chunk_id,
                'text': doc_text,
                'metadata': metadata,
                'source': metadata.get('source', 'unknown'),
                'chapter': metadata.get('chapter', ''),
                'section': metadata.get('section', '')
            }})
        except Exception as e:
            logger.error(f"Get citation detail failed: {e}")
            return jsonify({'success': False, 'error': {'code': 'SERVER_ERROR', 'message': str(e)}}), 500

    def get_available_models(self):
        """获取可用模型列表"""
        models = [
            {
                'id': 'minimax',
                'name': 'MiniMax',
                'model_name': Config.MODEL_NAME if hasattr(Config, 'MODEL_NAME') else 'abab6.5s-chat',
                'enabled': bool(getattr(Config, 'MINIMAX_API_KEY', '')),
                'provider': 'minimax'
            },
            {
                'id': 'openai',
                'name': 'OpenAI',
                'model_name': Config.OPENAI_MODEL_NAME if hasattr(Config, 'OPENAI_MODEL_NAME') else 'gpt-4',
                'enabled': bool(getattr(Config, 'OPENAI_API_KEY', '')),
                'provider': 'openai'
            }
        ]

        current_model = 'minimax' if Config.LLM_PROVIDER.lower() == 'minimax' else 'openai'

        return jsonify({'success': True, 'data': {
            'current_model': current_model,
            'available_models': models
        }})

    def switch_model(self):
        """切换当前使用的模型"""
        data = request.get_json(silent=True) or {}
        model_id = data.get('model_id', '')

        if model_id not in ('minimax', 'openai'):
            return jsonify({'success': False, 'error': {'code': 'INVALID_MODEL', 'message': '不支持的模型ID'}}), 400

        if model_id == 'minimax':
            if not getattr(Config, 'MINIMAX_API_KEY', ''):
                return jsonify({'success': False, 'error': {'code': 'NOT_CONFIGURED', 'message': 'MiniMax API Key未配置'}}), 400
        else:
            if not getattr(Config, 'OPENAI_API_KEY', ''):
                return jsonify({'success': False, 'error': {'code': 'NOT_CONFIGURED', 'message': 'OpenAI API Key未配置'}}), 400

        Config.LLM_PROVIDER = model_id.upper()

        return jsonify({'success': True, 'data': {
            'message': f'模型已切换至: {model_id}',
            'current_model': model_id
        }})
